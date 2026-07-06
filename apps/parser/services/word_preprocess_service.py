"""Word document preprocessing service.

Extract structured information from Word .docx files:
- Question numbers and their order
- Section titles
- Extracted images (diagrams, charts)
- Page structure information

This information is used to validate and enhance the AI parsing results.
"""
import os
import re
import logging
from pathlib import Path
from django.conf import settings

logger = logging.getLogger(__name__)

# Patterns for question numbers in Chinese exam papers
QUESTION_NO_PATTERNS = [
    r'^(?:第\s*)?([一二三四五六七八九十\d]+[、.\s]\d*)',   # "一、1" or "1."
    r'^([1-9]\d*)[、.\s]',                                    # "1." "1、"
    r'^（\s*([1-9]\d*)\s*）',                                 # "（1）"
    r'^\(\s*([1-9]\d*)\s*\)',                                 # "(1)"
]

# Keywords that suggest a question references an image
IMAGE_REF_KEYWORDS = [
    '如图', '见图', '观察下图', '下图所示', '坐标系',
    '统计图', '直方图', '折线图', '饼图', '散点图',
    '示意图', '几何图', '四棱锥', '三棱柱', '圆锥',
]

# Section title patterns
SECTION_PATTERN = re.compile(r'^([一二三四五六七八九十]+)[、.\s](.+?)(?:（|（|\(|$)')


class WordPreprocessor:
    """Extract structured information from a Word .docx file."""

    def __init__(self, docx_path: str):
        """
        Args:
            docx_path: Absolute path to the .docx file.
        """
        self.docx_path = docx_path
        self.paper_id = Path(docx_path).parent.parent.name
        self.question_numbers = []  # List of (sort_order, question_no, paragraph_text)
        self.section_titles = []    # List of (section_no, section_title)
        self.extracted_images = []  # List of dict: {filename, rel_path, type, index}
        self.full_text = ''         # Full document text

    def run(self) -> dict:
        """Run the full extraction pipeline.

        Returns:
            dict with keys:
                - question_numbers: list of (sort_order, question_no)
                - section_titles: list of (section_no, title_text)
                - images: list of image info dicts
                - image_refs: list of paragraphs that reference images
                - full_text: full document text
        """
        try:
            import docx
        except ImportError:
            logger.error('python-docx not installed')
            return self._empty_result()

        doc = docx.Document(self.docx_path)
        self._extract_full_text(doc)
        self._extract_sections(doc)
        self._extract_question_numbers(doc)
        self._extract_images(doc)
        self._find_image_references()

        return {
            'question_numbers': self.question_numbers,
            'section_titles': self.section_titles,
            'images': self.extracted_images,
            'image_refs': self.image_references,
            'full_text': self.full_text,
        }

    def _extract_full_text(self, doc):
        """Extract the full text content of the document."""
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        self.full_text = '\n'.join(paragraphs)

    def _extract_sections(self, doc):
        """Extract section titles like "一、选择题"."""
        for para in doc.paragraphs:
            text = para.text.strip()
            match = SECTION_PATTERN.match(text)
            if match:
                section_no = match.group(1)
                section_title = match.group(2).strip()
                self.section_titles.append((section_no, section_title))

    def _extract_question_numbers(self, doc):
        """Extract question numbers from the document."""
        sort_order = 0
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Skip short paragraphs (likely not questions)
            if len(text) < 5:
                continue

            for pattern in QUESTION_NO_PATTERNS:
                match = re.match(pattern, text)
                if match:
                    qno = match.group(1)
                    # Skip if this looks like a sub-question number
                    if len(qno) > 3:
                        continue
                    self.question_numbers.append((sort_order, qno, text[:100]))
                    sort_order += 1
                    break

    def _extract_images(self, doc):
        """Extract embedded images from the Word document."""
        try:
            import docx
            from docx.oxml.ns import qn
        except ImportError:
            return

        doc = docx.Document(self.docx_path)
        image_dir = settings.MEDIA_ROOT / 'exams' / str(self.paper_id) / 'word_images'
        image_dir.mkdir(parents=True, exist_ok=True)

        image_idx = 0
        # Iterate through all inline images in the document
        for rel in doc.part.rels.values():
            if 'image' in rel.reltype:
                image_idx += 1
                try:
                    # Get image data
                    image_data = rel.target_part.blob
                    content_type = rel.target_part.content_type

                    # Determine extension
                    if 'png' in content_type:
                        ext = '.png'
                    elif 'jpeg' in content_type or 'jpg' in content_type:
                        ext = '.jpg'
                    elif 'gif' in content_type:
                        ext = '.gif'
                    elif 'bmp' in content_type:
                        ext = '.bmp'
                    elif 'emf' in content_type:
                        ext = '.emf'
                    else:
                        ext = '.png'

                    filename = f'word_img_{image_idx:03d}{ext}'
                    filepath = image_dir / filename

                    with open(filepath, 'wb') as f:
                        f.write(image_data)

                    rel_path = os.path.relpath(filepath, settings.MEDIA_ROOT)
                    self.extracted_images.append({
                        'filename': filename,
                        'rel_path': rel_path,
                        'type': 'diagram',
                        'index': image_idx,
                        'size': len(image_data),
                    })
                    logger.info(f'Extracted image: {filename} ({len(image_data)} bytes)')
                except Exception as e:
                    logger.warning(f'Failed to extract image {image_idx}: {e}')

    def _find_image_references(self):
        """Find paragraphs that reference images."""
        self.image_references = []
        img_idx = 0
        for order, qno, text in self.question_numbers:
            if any(kw in text for kw in IMAGE_REF_KEYWORDS):
                self.image_references.append({
                    'question_no': qno,
                    'sort_order': order,
                    'text_preview': text[:100],
                    '_image_index': img_idx,
                })
                img_idx += 1

    def _empty_result(self):
        return {
            'question_numbers': [],
            'section_titles': [],
            'images': [],
            'image_refs': [],
            'full_text': '',
        }
