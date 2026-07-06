"""Root conftest.py — shared pytest fixtures for the front project."""
import io
import os
import pytest
from django.contrib.auth import get_user_model
from rest_framework.test import APIClient
from docx import Document

UserAccount = get_user_model()


# ─── User fixtures ───────────────────────────────────────────────

@pytest.fixture
def admin_user(db):
    """Create an admin user."""
    return UserAccount.objects.create(
        role_type='admin',
        mobile='13900000001',
        display_name='管理员',
        password='pbkdf2_sha256$dummy',
    )


@pytest.fixture
def teacher_user(db):
    """Create a teacher user."""
    return UserAccount.objects.create(
        role_type='teacher',
        mobile='13900000002',
        display_name='测试老师',
        subject='数学',
        stages=['初中', '高中'],
        password='pbkdf2_sha256$dummy',
    )


@pytest.fixture
def student_user(db):
    """Create a student user."""
    return UserAccount.objects.create(
        role_type='student',
        mobile='13900000003',
        display_name='测试学生',
        password='pbkdf2_sha256$dummy',
    )


# ─── API client fixtures ─────────────────────────────────────────

@pytest.fixture
def api_client():
    """Return an unauthenticated APIClient."""
    return APIClient()


@pytest.fixture
def admin_client(db, admin_user):
    """APIClient authenticated as admin."""
    client = APIClient()
    client.force_authenticate(user=admin_user)
    return client


@pytest.fixture
def teacher_client(db, teacher_user):
    """APIClient authenticated as teacher."""
    client = APIClient()
    client.force_authenticate(user=teacher_user)
    return client


@pytest.fixture
def student_client(db, student_user):
    """APIClient authenticated as student."""
    client = APIClient()
    client.force_authenticate(user=student_user)
    return client


# ─── File fixtures ───────────────────────────────────────────────

@pytest.fixture
def sample_docx_file(tmp_path):
    """Create a temporary .docx file for upload testing."""
    doc = Document()
    doc.add_heading('测试试卷', level=1)
    doc.add_paragraph('这是一份测试用的试卷文档。')
    doc.add_paragraph('第一题：选择题（5分）')
    doc.add_paragraph('以下哪个选项是正确的？')
    doc.add_paragraph('A. 选项一    B. 选项二    C. 选项三    D. 选项四')
    file_path = tmp_path / 'test-paper.docx'
    doc.save(str(file_path))
    return file_path


@pytest.fixture
def sample_pdf_file(tmp_path):
    """Create a minimal .pdf file for upload testing."""
    file_path = tmp_path / 'test-paper.pdf'
    # Minimal valid PDF (1.4)
    pdf_content = (
        b'%PDF-1.4\n'
        b'1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n'
        b'2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n'
        b'3 0 obj<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R'
        b'/Resources<<>>>>endobj\n'
        b'trailer<</Root 1 0 R>>\n'
        b'%%EOF'
    )
    file_path.write_bytes(pdf_content)
    return file_path


@pytest.fixture
def sample_doc_file(tmp_path):
    """Create a minimal .doc file (binary placeholder) for upload testing."""
    file_path = tmp_path / 'test-paper.doc'
    # Minimal OLE Compound File header (simplified — just needs to pass extension check)
    file_path.write_bytes(b'\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1' + b'\x00' * 100)
    return file_path


# ─── Institution / Class fixtures ────────────────────────────────

@pytest.fixture
def sample_institution(db, admin_user):
    """Create a sample institution."""
    from apps.institutions.models import Institution
    return Institution.objects.create(
        institution_name='测试学校',
        contact_name='联系人',
        contact_phone='13800000001',
        created_by=admin_user,
    )


@pytest.fixture
def sample_class(db, sample_institution, teacher_user):
    """Create a sample class with ClassTeacher record."""
    from apps.institutions.models import Class, InstitutionMember, ClassTeacher
    # Ensure teacher is member of institution
    InstitutionMember.objects.update_or_create(
        institution=sample_institution,
        user=teacher_user,
        defaults={'role': 'teacher'},
    )
    cls = Class.objects.create(
        institution=sample_institution,
        class_name='测试班级',
        creator_teacher=teacher_user,
        description='测试班级描述',
    )
    # Create ClassTeacher record for permission checks
    ClassTeacher.objects.create(
        class_obj=cls,
        teacher=teacher_user,
        role='owner',
    )
    return cls


# ─── Mission / Question fixtures ─────────────────────────────────

@pytest.fixture
def sample_mission(db, teacher_user):
    """Create a sample learning mission."""
    from apps.missions.models import LearningMission
    return LearningMission.objects.create(
        mission_name='测试任务',
        goal_text='测试学习目标',
        status='draft',
        creator_teacher_id=teacher_user,
    )


@pytest.fixture
def sample_mission_level(db, sample_mission):
    """Create a sample mission level."""
    from apps.missions.models import MissionLevel
    return MissionLevel.objects.create(
        mission=sample_mission,
        level_no=1,
        level_name='第一关',
        level_type='practice',
        mode_policy='block_a',
    )


@pytest.fixture
def sample_paper(db, teacher_user):
    """Create a sample exam paper."""
    from apps.papers.models import ExamPaper
    return ExamPaper.objects.create(
        title='测试试卷',
        subject='数学',
        stage='初中',
        grade='9',
        source_file_path='uploads/test-paper.docx',
        status='uploaded',
        uploaded_by=teacher_user,
    )


@pytest.fixture
def sample_question(db, sample_paper):
    """Create a sample exam question."""
    from apps.parser.models import ExamQuestion
    return ExamQuestion.objects.create(
        paper=sample_paper,
        question_no='1',
        question_type='single_choice',
        subject='数学',
        stem='以下哪个选项是正确的？',
        answer='A',
        difficulty=3.00,
    )
