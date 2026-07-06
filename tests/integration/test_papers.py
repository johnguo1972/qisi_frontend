"""Integration tests for paper upload and import batches."""
import pytest
from django.core.files.uploadedfile import SimpleUploadedFile
from apps.papers.models import ExamPaper, ParseTask


@pytest.mark.django_db
class TestPaperUpload:
    """Test paper upload via import_batch_list endpoint."""

    def _make_docx_file(self, name='test-paper.docx'):
        """Create a SimpleUploadedFile for a docx."""
        from docx import Document
        import io
        doc = Document()
        doc.add_heading('测试试卷', level=1)
        doc.add_paragraph('测试内容')
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return SimpleUploadedFile(name, buf.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    def _make_doc_file(self, name='test-paper.doc'):
        """Create a SimpleUploadedFile for a doc (binary placeholder)."""
        content = b'\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1' + b'\x00' * 100
        return SimpleUploadedFile(name, content, content_type='application/msword')

    def _make_pdf_file(self, name='test-paper.pdf'):
        """Create a SimpleUploadedFile for a minimal PDF."""
        content = (
            b'%PDF-1.4\n1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n'
            b'2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n'
            b'3 0 obj<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R>>endobj\n'
            b'trailer<</Root 1 0 R>>\n%%EOF'
        )
        return SimpleUploadedFile(name, content, content_type='application/pdf')

    def _make_txt_file(self, name='test-paper.txt'):
        """Create a SimpleUploadedFile for a txt (should be rejected)."""
        return SimpleUploadedFile(name, b'This is a text file', content_type='text/plain')

    def test_upload_docx_success(self, teacher_client):
        """Upload .docx should succeed."""
        f = self._make_docx_file()
        resp = teacher_client.post(
            '/api/v1/questions/import-batches',
            {'file': f},
            format='multipart'
        )
        assert resp.status_code == 200, f"Upload failed: {resp.json()}"
        data = resp.json()
        assert data['code'] == 0
        assert data['data']['paper_id'] is not None

    def test_upload_doc_success(self, teacher_client):
        """Upload .doc should succeed."""
        f = self._make_doc_file()
        resp = teacher_client.post(
            '/api/v1/questions/import-batches',
            {'file': f},
            format='multipart'
        )
        assert resp.status_code == 200, f"Upload failed: {resp.json()}"

    def test_upload_pdf_success(self, teacher_client):
        """Upload .pdf should succeed."""
        f = self._make_pdf_file()
        resp = teacher_client.post(
            '/api/v1/questions/import-batches',
            {'file': f},
            format='multipart'
        )
        assert resp.status_code == 200, f"Upload failed: {resp.json()}"

    def test_upload_txt_rejected(self, teacher_client):
        """Upload .txt should be rejected with 400."""
        f = self._make_txt_file()
        resp = teacher_client.post(
            '/api/v1/questions/import-batches',
            {'file': f},
            format='multipart'
        )
        assert resp.status_code == 400
        data = resp.json()
        assert '仅支持' in data['message']

    def test_upload_missing_file(self, teacher_client):
        """Upload without file should return 400."""
        resp = teacher_client.post(
            '/api/v1/questions/import-batches',
            {},
            format='multipart'
        )
        assert resp.status_code == 400

    def test_duplicate_paper_rejected(self, teacher_client):
        """Upload same filename twice should return 409 on second attempt."""
        f = self._make_docx_file('duplicate-test.docx')
        # First upload
        resp1 = teacher_client.post(
            '/api/v1/questions/import-batches',
            {'file': f},
            format='multipart'
        )
        assert resp1.status_code == 200

        # Second upload (same name)
        f2 = self._make_docx_file('duplicate-test.docx')
        resp2 = teacher_client.post(
            '/api/v1/questions/import-batches',
            {'file': f2},
            format='multipart'
        )
        assert resp2.status_code == 409
        assert '已存在' in resp2.json()['message']

    def test_import_batches_list(self, teacher_client):
        """List import batches should return current user's uploads."""
        # Upload first
        f1 = self._make_docx_file('batch-test-1.docx')
        teacher_client.post('/api/v1/questions/import-batches', {'file': f1}, format='multipart')

        # Upload second
        f2 = self._make_docx_file('batch-test-2.docx')
        teacher_client.post('/api/v1/questions/import-batches', {'file': f2}, format='multipart')

        # List
        resp = teacher_client.get('/api/v1/questions/import-batches')
        assert resp.status_code == 200
        data = resp.json()
        assert data['code'] == 0
        items = data['data']['items']
        assert len(items) >= 2

    def test_import_batches_isolation(self, teacher_client, student_user):
        """Import batches should only show current user's uploads."""
        # Teacher uploads
        f = self._make_docx_file('isolation-test.docx')
        teacher_client.post('/api/v1/questions/import-batches', {'file': f}, format='multipart')

        # Student lists — should not see teacher's upload
        from rest_framework.test import APIClient
        student_client = APIClient()
        student_client.force_authenticate(user=student_user)
        resp = student_client.get('/api/v1/questions/import-batches')
        assert resp.status_code == 200
        items = resp.json()['data']['items']
        # Should not include the teacher's paper
        for item in items:
            title = item.get('paper_title', '')
            assert title != 'isolation-test'

    def test_import_batches_empty(self, teacher_client):
        """Empty import batches should return empty list."""
        resp = teacher_client.get('/api/v1/questions/import-batches')
        assert resp.status_code == 200
        data = resp.json()
        items = data['data']['items']
        assert isinstance(items, list)
